from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
# 한국 사람이니까 폰트 크기 및 센치미터를 불러와서 사이즈를 적용하겠습니다. 인치도 있으니 필요시에 사용하세용  
from docx.shared import RGBColor
# 색깔을 위해 사용합니다. 


def make_blnk(doc,count):
    for i in range(count):
        doc.add_paragraph()
    

def image_tables(doc, images, front_gab=1, end_gab=1):
    '''
    docx의 Document 객체를 입력받아 앞 뒤 공백을 만들고, 
    이미지 정보가 들어있는 리스트를 받아 하나의 행에 모든 이미지를 배치합니다.
    
    Args:
        doc: Document 객체 (기존 문서에 추가)
        images: 이미지 정보 리스트 
                [{'path': '경로', 'width': Cm(8), 'caption': '설명'}, ...]
        front_gab: 앞쪽 빈 줄 개수
        end_gab: 뒤쪽 빈 줄 개수
    '''
    # 앞쪽 공백 추가
    if front_gab:
        make_blnk(doc, front_gab)
    
    # 이미지 테이블 생성
    image_table = doc.add_table(rows=1, cols=len(images))
    
    # 각 이미지를 셀에 배치
    for idx, image_info in enumerate(images):
        cell = image_table.cell(0, idx)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 이미지 추가
        run = paragraph.add_run()
        run.add_picture(
            image_info['path'], 
            width=image_info.get('width', Cm(8))
        )
    
    # 캡션 테이블 생성 (필요한 경우)
    if any('caption' in img for img in images):
        caption_table = doc.add_table(rows=1, cols=len(images))
        for idx, image_info in enumerate(images):
            if 'caption' in image_info:
                cell = caption_table.cell(0, idx)
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run(image_info['caption'])
                run.italic = True
    
    # 뒤쪽 공백 추가
    if end_gab:
        make_blnk(doc, end_gab)

