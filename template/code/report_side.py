from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
# 한국 사람이니까 폰트 크기 및 센치미터를 불러와서 사이즈를 적용하겠습니다. 인치도 있으니 필요시에 사용하세용  
from docx.shared import RGBColor
# 색깔을 위해 사용합니다. 
from datetime import date

documnet = Document()
# 객체 선언 

logos_and_date = documnet.add_table(rows=1,cols=3)
logo = logos_and_date.cell(0,0)
logo_section = logo.paragraphs[0]
logo_section.alignment = WD_ALIGN_PARAGRAPH.LEFT
logo_image = logo_section.add_run().add_picture('../../image/roa.png')
logo_image.height = Cm(1)
logo_image.width = Cm(1)

# 로고 삽입 
dates = logos_and_date.cell(0,2)
dates_section = dates.paragraphs[0]
dates_section.alignment = WD_ALIGN_PARAGRAPH.RIGHT
date_insert = dates_section.add_run(f'작성일 : {date.today()}')
date_insert.font.size = (Pt(10))
date_insert.font.color.rgb = RGBColor(128,128,128)

title = documnet.add_heading('보고서 입니다.', level = 0)

title.alignment =WD_ALIGN_PARAGRAPH.CENTER
# 중앙 정렬을 지정합니다. 매개변수로 right , left도 있습니다 

documnet.add_paragraph()
# 이렇게 하면 빈줄을 추가합니다. 
sunmmer=documnet.add_heading('요약',level=1)
sunmmer.alignment = WD_ALIGN_PARAGRAPH.CENTER
# 중앙정렬
chorock = documnet.add_paragraph()
# 줄을 잡아두고 
chorock.add_run('요약을 할 예정입니다\n').bold = True 
chorock.add_run('메이플을 해보셨나요? 오늘 집가서 할 예정이긴 한데 요즘 좀 재미없네요\n').font.size = Pt(11)
chorock.add_run('거참 딱 접기 좋은 날씨네').font.color.rgb = RGBColor(128, 128, 128)
chorock.alignment = WD_ALIGN_PARAGRAPH.CENTER
# 2개 넣을 거니까 1행 2열이면 되겠죠? 
image_table = documnet.add_table(rows=1,cols=2)
left_cell = image_table.cell(0,0)
left_cell_section = left_cell.paragraphs[0]
left_cell_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
left_cell_section.add_run().add_picture(r'../../image/mul.png',width = Cm(8))

right_cell = image_table.cell(0,1)
right_cell_section = right_cell.paragraphs[0]
right_cell_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
right_cell_section.add_run().add_picture(r'../../image/mul2.png',width = Cm(8))
text_table = documnet.add_table(cols=2,rows=1)

ltext_cell = text_table.cell(0,0)
ltext_cell_section = ltext_cell.paragraphs[0]
ltext_cell_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
# 위치 같아야 그 위치 밑에 텍스트가 박힙니다. 
ltext_cell_section.add_run('수영물개')

rtext_cell = text_table.cell(0,1)
rtext_cell_section = rtext_cell.paragraphs[0]
rtext_cell_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
rtext_cell_section.add_run('뒹굴물개')
tables=documnet.add_table(2,2)
tables.style = 'Table Grid'

for row in tables.rows:
    for cell in row.cells:
        cell.text = '이미지는 전부 제미나이로 생성했습니다.'

documnet.save('../../learn_word/test_document.docx')